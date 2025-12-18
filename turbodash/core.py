import numpy as np
from scipy.optimize import root_scalar

import jaxprop as jxp


def compute_performance_stage(
    stator_inlet_angle,
    stator_exit_angle,
    degree_reaction,
    blade_velocity_ratio,
    radius_ratio_23=1.00,
    radius_ratio_34=1.00,
    loss_coeff_stator=0.0,
    loss_coeff_rotor=0.0,
):
    r"""
    Compute flow and performance parameters of a turbine stage.

    The function evaluates the dimensionless flow and loading coefficients,
    outlet flow angles, and total-to-static and total-to-total efficiencies
    for given stage geometry and velocity-triangle parameters.

    Angles are defined in the tangential-axial velocity plane, positive in the direction of rotation

    The calculation of the velocity triangles assumes isentropic flow.
    The calculation of the losses is decoupled from the velocity triangles.

    Parameters
    ----------
    stator_inlet_angle : float
        Absolute flow angle at stator inlet (α₁) in degrees.
    stator_exit_angle : float
        Absolute flow angle at stator exit (α₂) in degrees.
    degree_reaction : float
        Degree of reaction, defined as the static enthalpy drop in the rotor
        over the total stage enthalpy drop (R = Δh_rotor / Δh_stage).
    blade_velocity_ratio : float
        Blade speed ratio (nu = U / √(2Δh_0s)), a non-dimensional velocity ratio.
    radius_ratio_23 : float, optional
        Mean-to-outlet radius ratio (rho = r_2 / r_3). Default is 1.0.
    radius_ratio_34 : float, optional
        Mean-to-outlet radius ratio (rho = r_3 / r_4). Default is 1.0.
    loss_coeff_stator : float, optional
        Stator loss coefficient, ζ_stator = (Δh_loss / ½v²). Default is 0.0.
    loss_coeff_rotor : float, optional
        Rotor loss coefficient, ζ_rotor = (Δh_loss / ½w²). Default is 0.0.

    """

    # Rename variables
    R = degree_reaction
    nu = blade_velocity_ratio
    alpha1 = np.deg2rad(stator_inlet_angle)
    alpha2 = np.deg2rad(stator_exit_angle)
    tan_alpha1 = np.tan(alpha1)
    tan_alpha2 = np.tan(alpha2)

    # Compute angle after the interspace
    tan_alpha3 = np.tan(alpha2) * radius_ratio_23**-1
    alpha3 = np.arctan(tan_alpha3)

    # φ² = (1 - R) / [ν² * ((1 + tan²α₂) - R(1 + tan²α₁))]
    phi_sq = (1 - R) / (nu**2 * ((1 + tan_alpha3**2) - R * (1 + tan_alpha1**2)))
    phi = np.sqrt(phi_sq)

    # tanβ₂ = tanα₂ - ρ/φ
    tan_beta3 = tan_alpha3 - radius_ratio_34 / phi
    beta3 = np.arctan(tan_beta3)

    # tan²β₃ = tan²β₂ - tan²α₂ - 1 + 1/(ν²φ²) + (1-ρ²)/φ²
    # Use the negative root of this second order equation
    tan_beta4_sq = (
        tan_beta3**2
        - tan_alpha3**2
        - 1
        + 1 / (nu**2 * phi**2)
        + (1 - radius_ratio_34**2) / phi**2
    )
    tan_beta4 = -np.sqrt(np.clip(tan_beta4_sq, 0, None))
    beta4 = np.arctan(tan_beta4)

    # tanα₃ = tanβ₃ + 1/φ
    tan_alpha4 = tan_beta4 + 1 / phi
    alpha4 = np.arctan(tan_alpha4)

    # Both expressions for the work coefficient give the same result
    # ψ = φ (ρ tanα₂ - tanα₃)
    # ψ = (1/(2ν²)) [1 - ν² φ² (1 + tan²α₃)]
    psi = phi * (radius_ratio_34 * tan_alpha3 - tan_alpha4) + 1e-12
    # psi2 = 0.5 / (nu**2) * (1 - (nu**2) * phi**2 * (1 + tan_alpha4**2))

    # Compute losses and efficiency in a decoupled way
    d_eta_ke = nu**2 * phi**2 * (1 + tan_alpha4**2)
    loss_stator = (1 + tan_alpha2**2) * loss_coeff_stator
    loss_rotor = (1 + tan_beta4**2) * loss_coeff_rotor
    eta_tt = psi / (psi + 0.5 * phi**2 * (loss_stator + loss_rotor))
    eta_ts = (1 / eta_tt + 0.5 * phi**2 / psi * (1 + tan_alpha4**2)) ** -1

    return {
        "phi": phi,
        "psi": psi,
        "alpha3": np.rad2deg(alpha3),
        "alpha4": np.rad2deg(alpha4),
        "beta3": np.rad2deg(beta3),
        "beta4": np.rad2deg(beta4),
        "eta_ts": eta_ts,
        "eta_tt": eta_tt,
        "d_eta_ke": d_eta_ke,
    }


def compute_performance_repeating_stage(alpha2_deg, R, nu):
    """
    Solve for α₁ such that α₁ = α₃ (repeating stage) using a Newton solver,
    and compute flow/performance quantities.

    Supports arbitrary array shapes for alpha2_deg, R, and nu (NumPy broadcasting).

    Parameters
    ----------
    alpha2_deg : float or array_like
        Absolute flow angle at rotor inlet [degrees].
    R : float or array_like
        Degree of reaction.
    nu : float or array_like
        Blade-to-spouting velocity ratio nu = u / v₀.

    Returns
    -------
    results : dict[str, np.ndarray]
        Arrays for α₁, α₃, β₂, β₃, φ, ψ, and η_ts with broadcasted shape (trivial dims squeezed).
    """

    # --- broadcast all input arrays ---
    alpha2_deg, R, nu = np.broadcast_arrays(alpha2_deg, R, nu)
    shape = alpha2_deg.shape

    # --- preallocate outputs ---
    alpha1_out = np.full(shape, np.nan)
    alpha3_out = np.full(shape, np.nan)
    alpha4_out = np.full(shape, np.nan)
    beta3_out = np.full(shape, np.nan)
    beta4_out = np.full(shape, np.nan)
    phi_out = np.full(shape, np.nan)
    psi_out = np.full(shape, np.nan)
    eta_ts_out = np.full(shape, np.nan)

    # --- iterate over all input combinations ---
    it = np.nditer([alpha2_deg, R, nu], flags=["multi_index"])

    for a2, r, n in it:
        idx = it.multi_index
        alpha2_deg = float(a2)
        R_val = float(r)
        nu_val = float(n)

        def closure(alpha1_deg):
            res = compute_performance_stage(alpha1_deg, alpha2_deg, R_val, nu_val)
            residual = alpha1_deg - res["alpha4"]
            return residual

        # --- use Newton solver starting from α₁ = 0° ---
        sol = root_scalar(closure, x0=0, method="newton", maxiter=50)
        if not sol.converged:
            print(sol)
            raise RuntimeError
        alpha1_deg = sol.root

        # recompute full stage quantities
        res = compute_performance_stage(alpha1_deg, alpha2_deg, R_val, nu_val)

        # store results
        phi_out[idx] = res["phi"]
        psi_out[idx] = res["psi"]
        alpha1_out[idx] = alpha1_deg
        alpha3_out[idx] = res["alpha3"]
        alpha4_out[idx] = res["alpha4"]
        beta3_out[idx] = res["beta3"]
        beta4_out[idx] = res["beta4"]
        eta_ts_out[idx] = res["eta_ts"]

    # --- return all results (squeezed to remove trivial dims) ---
    return {
        "phi": np.squeeze(phi_out),
        "psi": np.squeeze(psi_out),
        "alpha1": np.squeeze(alpha1_out),
        "alpha3": np.squeeze(alpha3_out),
        "beta2": np.squeeze(beta3_out),
        "beta3": np.squeeze(beta4_out),
        "eta_ts": np.squeeze(eta_ts_out),
    }


def assert_velocity_triangle(v, w, u, alpha_deg, beta_deg, label, rtol=1e-8):
    lhs = v * np.sin(np.deg2rad(alpha_deg))
    rhs = w * np.sin(np.deg2rad(beta_deg)) + u

    assert np.isfinite(lhs) and np.isfinite(rhs), (
        f"{label}: non-finite velocity triangle terms " f"(lhs={lhs}, rhs={rhs})"
    )

    assert np.isclose(lhs, rhs, rtol=rtol, atol=1e-5), (
        f"{label}: velocity triangle not closed\n"
        f"  v*sin(alpha) = {lhs}\n"
        f"  w*sin(beta)+u = {rhs}\n"
        f"  alpha = {alpha_deg} deg, beta = {beta_deg} deg, u = {u}"
    )


def compute_stage_meanline(
    fluid_name,
    inlet_property_pair_string,
    inlet_property_1,
    inlet_property_2,
    exit_pressure,
    mass_flow_rate,
    stator_inlet_angle,
    stator_exit_angle,
    blade_velocity_ratio,
    degree_reaction,
    radius_ratio_12,
    radius_ratio_23,
    radius_ratio_34,
    height_radius_ratio,
    zweiffel_stator,
    zweiffel_rotor,
    loss_coeff_stator,
    loss_coeff_rotor,
    stage_type="radial",
):
    """
    Meanline turbine stage model with geometry sizing.

    - Velocity triangles and efficiencies from compute_performance_stage
    - Constant meridional velocity
    - Isentropic state reconstruction using (h, s)
    - Geometry from continuity + dimensionless parameters
    - Blade count from Zweifel criterion
    """

    # ------------------------------------------------------------------
    # Stage performance (velocity triangles, efficiencies)
    # ------------------------------------------------------------------
    RR_12 = radius_ratio_12
    RR_23 = radius_ratio_23
    RR_34 = radius_ratio_34
    perf = compute_performance_stage(
        stator_inlet_angle=stator_inlet_angle,
        stator_exit_angle=stator_exit_angle,
        degree_reaction=degree_reaction,
        blade_velocity_ratio=blade_velocity_ratio,
        radius_ratio_34=RR_34,
        loss_coeff_stator=loss_coeff_stator,
        loss_coeff_rotor=loss_coeff_rotor,
    )

    # ------------------------------------------------------------------
    # Isentropic outlet enthalpy and spouting velocity
    # ------------------------------------------------------------------
    fluid = jxp.Fluid(fluid_name, backend="HEOS")
    ip = jxp.INPUT_PAIRS[inlet_property_pair_string]
    state_01 = fluid.get_state(ip, inlet_property_1, inlet_property_2)
    state_4s = fluid.get_state(jxp.PSmass_INPUTS, exit_pressure, state_01.s)
    h_01 = state_01.h
    h_4s = state_4s.h
    assert h_01 > h_4s, "Non-positive isentropic enthalpy drop"
    v0 = np.sqrt(2.0 * (h_01 - h_4s))

    # ------------------------------------------------------------------
    # Velocities
    # ------------------------------------------------------------------

    # Blade velocities
    phi = perf["phi"]
    nu = blade_velocity_ratio
    U = nu * v0
    u_1 = 0.0
    u_2 = 0.0
    u_3 = U * RR_34
    u_4 = U

    # Absolute velocities
    vm = phi * U
    alpha_1 = stator_inlet_angle
    alpha_2 = stator_exit_angle
    # tan_alpha_3 = np.tan(np.deg2rad(alpha_2)) * radius_ratio_23**-1
    # alpha_3 = np.rad2deg(np.atan2(tan_alpha_3, 1))
    alpha_3 = perf["alpha3"]
    alpha_4 = perf["alpha4"]
    v_1 = vm / np.cos(np.deg2rad(alpha_1))
    v_2 = vm / np.cos(np.deg2rad(alpha_2))
    v_3 = vm / np.cos(np.deg2rad(alpha_3))
    v_4 = vm / np.cos(np.deg2rad(alpha_4))

    # Relative velocities
    beta_1 = alpha_1
    beta_2 = alpha_2
    beta_3 = perf["beta3"]
    beta_4 = perf["beta4"]
    w_1 = vm / np.cos(np.deg2rad(beta_1))
    w_2 = vm / np.cos(np.deg2rad(beta_2))
    w_3 = vm / np.cos(np.deg2rad(beta_3))
    w_4 = vm / np.cos(np.deg2rad(beta_4))

    # Check there are no errrors in the implementation
    assert_velocity_triangle(v_1, w_1, u_1, alpha_1, beta_1, "Station 1")
    assert_velocity_triangle(v_2, w_2, u_2, alpha_2, beta_2, "Station 2")
    assert_velocity_triangle(v_3, w_3, u_3, alpha_3, beta_3, "Station 3")
    assert_velocity_triangle(v_4, w_4, u_4, alpha_4, beta_4, "Station 4")

    # ------------------------------------------------------------------
    # Static enthalpies
    # ------------------------------------------------------------------
    tan2_a1 = np.tan(np.deg2rad(alpha_1)) ** 2
    tan2_a2 = np.tan(np.deg2rad(alpha_2)) ** 2
    tan2_a3 = np.tan(np.deg2rad(alpha_3)) ** 2
    tan2_b3 = np.tan(np.deg2rad(beta_3)) ** 2
    tan2_b4 = np.tan(np.deg2rad(beta_4)) ** 2
    h_1 = state_01.h - 0.5 * v_1**2
    h_2 = h_1 - 0.5 * U**2 * phi**2 * (tan2_a2 - tan2_a1)
    h_3 = h_2 - 0.5 * U**2 * phi**2 * (tan2_a3 - tan2_a2)
    h_4 = h_3 - U**2 * (0.5 * phi**2 * (tan2_b4 - tan2_b3) - 0.5 * (1.0 - RR_34**2))

    # ------------------------------------------------------------------
    # EOS states (isentropic, using (h, s_01))
    # ------------------------------------------------------------------
    # Station states (1..4): (h_i, s_01)
    state_1 = fluid.get_state(jxp.HmassSmass_INPUTS, h_1, state_01.s)
    state_2 = fluid.get_state(jxp.HmassSmass_INPUTS, h_2, state_01.s)
    state_3 = fluid.get_state(jxp.HmassSmass_INPUTS, h_3, state_01.s)
    state_4 = fluid.get_state(jxp.HmassSmass_INPUTS, h_4, state_01.s)

    # Convenience aliases
    d_1, d_2, d_3, d_4 = state_1.d, state_2.d, state_3.d, state_4.d
    p_1, p_2, p_3, p_4 = state_1.p, state_2.p, state_3.p, state_4.p
    a_1, a_2, a_3, a_4 = state_1.a, state_2.a, state_3.a, state_4.a
    s_1, s_2, s_3, s_4 = state_1.s, state_2.s, state_3.s, state_4.s
    T_1, T_2, T_3, T_4 = state_1.T, state_2.T, state_3.T, state_4.T

    # Compute Mach numbers
    Ma_1, Ma_2, Ma_3, Ma_4 = w_1 / a_1, w_2 / a_2, w_3 / a_3, w_4 / a_4

    # ------------------------------------------------------------------
    # Geometry from continuity
    # ------------------------------------------------------------------
    # Compute radii
    r_1 = np.sqrt(mass_flow_rate / (2.0 * np.pi * d_1 * vm * height_radius_ratio))
    r_2 = r_1 / RR_12
    r_3 = r_2 / RR_23
    r_4 = r_3 / RR_34

    # Compute blade heights
    H_1 = mass_flow_rate / (2.0 * np.pi * r_1 * d_1 * vm)
    H_2 = mass_flow_rate / (2.0 * np.pi * r_2 * d_2 * vm)
    H_3 = mass_flow_rate / (2.0 * np.pi * r_3 * d_3 * vm)
    H_4 = mass_flow_rate / (2.0 * np.pi * r_4 * d_4 * vm)

    # Asset inlet height-to-radius ratio
    assert np.isclose(
        height_radius_ratio,
        H_1 / r_1,
        rtol=1e-6,
        atol=0.0,
    ), (
        f"Height-to-radius ratio mismatch: "
        f"target={height_radius_ratio}, computed={H_1 / r_1}"
    )

    # ------------------------------------------------------------------
    # Angular speed from blade speed and radius
    # ------------------------------------------------------------------
    omega_3 = u_3 / r_3
    omega_4 = u_4 / r_4
    RPM = omega_4 * 60 / (2 * np.pi)
    ws = (
        omega_3
        * (mass_flow_rate / state_4s.d) ** (1 / 2)
        * (state_01.h - state_4s.h) ** (-3 / 4)
    )
    assert np.isclose(
        omega_3, omega_4, rtol=1e-6
    ), f"Inconsistent rotational speed: omega_3={omega_3:0.2f}, omega_4={omega_4:0.2f}"

    # ------------------------------------------------------------------
    # Stator and rotor geometry
    # ------------------------------------------------------------------

    # Stator
    stator_geom = compute_blade_row_geometry(
        stage_type=stage_type,
        r_in=r_1,
        r_out=r_2,
        H_in=H_1,
        H_out=H_2,
        angle_in_deg=alpha_1,
        angle_out_deg=alpha_2,
        zweiffel=zweiffel_stator,
    )

    # Rotor
    rotor_geom = compute_blade_row_geometry(
        stage_type=stage_type,
        r_in=r_3,
        r_out=r_4,
        H_in=H_3,
        H_out=H_4,
        angle_in_deg=beta_3,
        angle_out_deg=beta_4,
        zweiffel=zweiffel_rotor,
    )

    # ------------------------------------------------------------------
    # Export results as nested dictionary
    # ------------------------------------------------------------------
    # Initialize container
    out = {}

    # Input variables
    out["inputs"] = {
        "fluid": fluid_name,
        "stage_type": stage_type,
        "inlet_property_pair": inlet_property_pair_string,
        "inlet_property_1": inlet_property_1,
        "inlet_property_2": inlet_property_2,
        "exit_pressure": exit_pressure,
        "mass_flow_rate": mass_flow_rate,
        "stator_inlet_angle": stator_inlet_angle,
        "stator_exit_angle": stator_exit_angle,
        "blade_velocity_ratio": blade_velocity_ratio,
        "degree_reaction": degree_reaction,
        "radius_ratio_12": RR_12,
        "radius_ratio_23": RR_23,
        "radius_ratio_34": RR_34,
        "height_radius_ratio": height_radius_ratio,
        "zweiffel_stator": zweiffel_stator,
        "zweiffel_rotor": zweiffel_rotor,
        "loss_coeff_stator": loss_coeff_stator,
        "loss_coeff_rotor": loss_coeff_rotor,
    }

    # Performance metrics
    out["stage_performance"] = {
        "efficiency_tt": perf["eta_tt"],
        "efficiency_ts": perf["eta_ts"],
        "pressure_ratio_ts": state_01.p / state_4s.p,
        "volume_ratio_ts": state_01.d / state_4s.d,
        "flow_coefficient": perf["phi"],
        "work_coefficient": perf["psi"],
        "degree_reaction": degree_reaction,
        "blade_velocity_ratio": blade_velocity_ratio,
        "specific_speed": ws,
        "spouting_velocity": v0,
        "rotor_exit_velocity": U,
        "rotational_speed": RPM,
        "mass_flow_rate": mass_flow_rate,
        "power_isentropic": mass_flow_rate * (state_01.h - h_4s),
        "power_actual_tt": mass_flow_rate * (state_01.h - h_4s) * perf["eta_tt"],
        "power_actual_ts": mass_flow_rate * (state_01.h - h_4s) * perf["eta_ts"],
    }

    # Cascade geometry
    out["geometry"] = {
        "stator": stator_geom,
        "rotor": rotor_geom,
    }

    # Thermodynamics and kinematics at each flow station
    out["flow_stations"] = [
        # Station 0 (inlet total state)
        {
            "p": state_01.p,
            "T": state_01.T,
            "d": state_01.d,
            "q": state_01.quality_mass,
            "Z": state_01.Z,
            "a": state_01.a,
            "h": state_01.h,
            "s": state_01.s,
            "v": 0.0,
            "w": 0.0,
            "u": 0.0,
            "alpha": alpha_1,
            "beta": beta_1,
            "Ma": 0.0,
            "r": r_1,
            "H": H_1,
        },
        # Station 1 (stator inlet)
        {
            "p": p_1,
            "T": T_1,
            "d": d_1,
            "q": state_1.quality_mass,
            "Z": state_1.Z,
            "a": a_1,
            "h": h_1,
            "s": s_1,
            "v": v_1,
            "w": w_1,
            "u": u_1,
            "alpha": alpha_1,
            "beta": beta_1,
            "Ma": Ma_1,
            "r": r_1,
            "H": H_1,
        },
        # Station 2 (stator exit / rotor inlet)
        {
            "p": p_2,
            "T": T_2,
            "d": d_2,
            "q": state_2.quality_mass,
            "Z": state_2.Z,
            "a": a_2,
            "h": h_2,
            "s": s_2,
            "v": v_2,
            "w": w_2,
            "u": u_2,
            "alpha": alpha_2,
            "beta": beta_2,
            "Ma": Ma_2,
            "r": r_2,
            "H": H_2,
        },
        # Station 3 (rotor exit, relative frame)
        {
            "p": p_3,
            "T": T_3,
            "d": d_3,
            "q": state_3.quality_mass,
            "Z": state_3.Z,
            "a": a_3,
            "h": h_3,
            "s": s_3,
            "v": v_3,
            "w": w_3,
            "u": u_3,
            "alpha": alpha_3,
            "beta": beta_3,
            "Ma": Ma_3,
            "r": r_3,
            "H": H_3,
        },
        # Station 4 (rotor exit / stage outlet)
        {
            "p": p_4,
            "T": T_4,
            "d": d_4,
            "q": state_4.quality_mass,
            "Z": state_4.Z,
            "a": a_4,
            "h": h_4,
            "s": s_4,
            "v": v_4,
            "w": w_4,
            "u": u_4,
            "alpha": alpha_4,
            "beta": beta_4,
            "Ma": Ma_4,
            "r": r_4,
            "H": H_4,
        },
    ]

    return out


def compute_blade_row_geometry(
    stage_type,
    r_in,
    r_out,
    H_in,
    H_out,
    angle_in_deg,
    angle_out_deg,
    zweiffel,
    maximum_thickness_to_chord=0.3,
    maximum_thickness_location=0.25,
    leading_edge_radius_to_max_thickness=0.50,
    trailing_edge_thickness_to_opening=0.05,
    trailing_edge_wedge_angle=10.0,
):
    """
    Compute blade chord, spacing, blade count, opening, and solidity
    using Zweifel criterion and return full geometry dictionary

    Parameters
    ----------
    stage_type : {"radial", "axial"}
    r_in, r_out : float
        Inner and outer radii of the blade row
    H_in, H_out : float
        Channel heights at inlet and outlet
    angle_in_deg, angle_out_deg : float
        Flow angles (absolute for stator, relative for rotor)
    zweiffel : float
        Zweifel loading coefficient
    """

    # Chord definition
    if stage_type == "radial":
        meridional_chord = r_out - r_in
    elif stage_type == "axial":
        meridional_chord = 0.75 * 0.5 * (H_in + H_out)
    else:
        raise ValueError(f"Invalid stage type: {stage_type}")

    # Integer number of blades from Zweifel criterion
    angle_in = np.deg2rad(angle_in_deg)
    angle_out = np.deg2rad(angle_out_deg)
    s_mean = (
        0.5
        * zweiffel
        * meridional_chord
        / (np.cos(angle_out) ** 2 * np.abs(np.tan(angle_in) - np.tan(angle_out)))
    )

    N_blades = int(np.ceil(np.pi * (r_in + r_out) / s_mean))
    s_mean = (np.pi * (r_in + r_out)) / N_blades
    solidity = meridional_chord / s_mean

    # Opening (cosine rule)
    s_out = 2 * np.pi * r_out / N_blades
    if stage_type == "radial":
        o = s_out * np.cos(np.abs(angle_out) - 0.5 * (2.0 * np.pi / N_blades))
    elif stage_type == "axial":
        o = s_out * np.cos(angle_out)
    else:
        raise ValueError(f"Invalid stage type: {stage_type}")

    # Flaring angle
    height = 0.5 * (H_in + H_out)
    aspect_ratio = height / meridional_chord
    flaring_angle = np.rad2deg(np.arctan(0.5 * (H_out - H_in) / meridional_chord))

    # Miscellaneous quantities for plotting
    maximum_thickness = meridional_chord * maximum_thickness_to_chord
    leading_edge_radius = maximum_thickness * leading_edge_radius_to_max_thickness
    trailing_edge_thickness = o * trailing_edge_thickness_to_opening

    # Return complete dictionary
    return dict(
        blade_count=N_blades,
        radius_in=r_in,
        radius_out=r_out,
        height=height,
        chord=meridional_chord,
        spacing=s_mean,
        opening=o,
        solidity=solidity,
        aspect_ratio=aspect_ratio,
        flaring_angle=flaring_angle,
        maximum_thickness=maximum_thickness,
        maximum_thickness_location=maximum_thickness_location,
        leading_edge_radius=leading_edge_radius,
        trailing_edge_thickness=trailing_edge_thickness,
        trailing_edge_wedge_angle=trailing_edge_wedge_angle,
        metal_angle_in=angle_in_deg,
        metal_angle_out=angle_out_deg,
    )


def _fmt(value, unit="-", width=10, prec=4):
    if value is None:
        return " " * width

    # Strings: right-aligned, same width as numbers, no unit
    if isinstance(value, str):
        return f"{value:>{width}s}"

    # Integers
    if isinstance(value, int):
        return f"{value:{width}d} {unit}".rstrip()

    # Floats
    if isinstance(value, float):
        return f"{value:{width}.{prec}f} {unit}".rstrip()

    # Fallback
    return f"{str(value):>{width}s}"


def _fmt_mm(value_m, width=10, prec=2):
    return _fmt(1e3 * value_m, "mm", width, prec)


def _fmt_deg(value, width=10, prec=2):
    return _fmt(value, "deg", width, prec)


def _section(title):
    print(f"\n{title}")
    print("=" * len(title))


def _pretty_name(key: str) -> str:
    return key.replace("_", " ").capitalize()


def _fmt_geom(key, value):
    """
    Geometry-aware formatter based on variable semantics.
    """
    if value is None:
        return " " * 12

    if key in {
        "radius_in",
        "radius_out",
        "height",
        "chord",
        "spacing",
        "opening",
        "maximum_thickness",
        "leading_edge_radius",
        "trailing_edge_thickness",
    }:
        return _fmt_mm(value)

    if key in {
        "flaring_angle",
        "trailing_edge_wedge_angle",
        "metal_angle_in",
        "metal_angle_out",
    }:
        return _fmt_deg(value)

    if key in {
        "aspect_ratio",
        "solidity",
        "flaring_angle",
        "trailing_edge_wedge_angle",
        "maximum_thickness_location",
    }:
        return _fmt(value, "-")

    if key == "blade_count":
        return _fmt(value, "-", prec=0)

    # Fallback
    return _fmt(value, "-")


def _print_geometry_block(title, geom):
    print(f"\n{title}:")
    for key, value in geom.items():
        label = _pretty_name(key)
        print(f"  {label:28s}: {_fmt_geom(key, value)}")


def print_stage(out):
    """
    ASCII-only, human-readable pretty print of turbine stage results.
    Assumes a fully nested and structured `out` dictionary.
    """

    inputs = out["inputs"]
    perf = out["stage_performance"]
    geom = out["geometry"]
    stations = out["flow_stations"]

    # ==============================================================
    # Boundary conditions and operating point
    # ==============================================================
    _section("Boundary conditions and operating point")

    print(f"{'Fluid':30s}: {_fmt(inputs['fluid'])}")
    print(f"{'Stage type':30s}: {_fmt(inputs['stage_type'])}")
    print(f"{'Inlet pressure':30s}: {_fmt(stations[0]['p'] / 1e5, 'bar')}")
    print(f"{'Exit pressure':30s}: {_fmt(stations[-1]['p'] / 1e5, 'bar')}")
    print(f"{'Mass flow rate':30s}: {_fmt(inputs['mass_flow_rate'], 'kg/s')}")
    print(f"{'Rotational speed':30s}: {_fmt(perf['rotational_speed'], 'rpm', prec=1)}")

    # ==============================================================
    # Stage performance
    # ==============================================================
    _section("Stage performance")

    print(f"{'Total-to-total efficiency':30s}: {_fmt(perf['efficiency_tt'], '-')}")
    print(f"{'Total-to-static efficiency':30s}: {_fmt(perf['efficiency_ts'], '-')}")
    print(
        f"{'Total-to-static pressure ratio':30s}: {_fmt(perf['pressure_ratio_ts'], '-')}"
    )
    print(f"{'Total-to-static volume ratio':30s}: {_fmt(perf['volume_ratio_ts'], '-')}")
    print(f"{'Flow coefficient':30s}: {_fmt(perf['flow_coefficient'], '-')}")
    print(f"{'Loading coefficient':30s}: {_fmt(perf['work_coefficient'], '-')}")
    print(f"{'Degree of reaction':30s}: {_fmt(perf['degree_reaction'], '-')}")
    print(f"{'Blade velocity ratio':30s}: {_fmt(perf['blade_velocity_ratio'], '-')}")
    print(f"{'Specific speed':30s}: {_fmt(perf['specific_speed'], '-')}")
    print(f"{'Spouting velocity':30s}: {_fmt(perf['spouting_velocity'], 'm/s')}")
    print(f"{'Rotor exit blade speed':30s}: {_fmt(perf['rotor_exit_velocity'], 'm/s')}")

    # ==============================================================
    # Geometry summary
    # ==============================================================
    _section("Geometry summary")
    _print_geometry_block("Stator", geom["stator"])
    _print_geometry_block("Rotor", geom["rotor"])

    # ==============================================================
    # Flow stations
    # ==============================================================
    _section("Flow stations")
    header = (
        f"{'Stn':>3s} "
        f"{'p [bar]':>10s} "
        f"{'T [K]':>9s} "
        f"{'rho [kg/m3]':>12s} "
        f"{'v [m/s]':>10s} "
        f"{'w [m/s]':>10s} "
        f"{'Ma':>7s} "
        f"{'alpha [deg]':>12s} "
        f"{'beta [deg]':>11s} "
        f"{'r [mm]':>9s} "
        f"{'H [mm]':>9s}"
    )
    print()
    print("-" * len(header))

    print(header)

    for i, st in enumerate(stations):
        print(
            f"{i:3d} "
            f"{st['p']/1e5:10.3f} "
            f"{st['T']:9.2f} "
            f"{st['d']:12.4f} "
            f"{st['v']:10.2f} "
            f"{st['w']:10.2f} "
            f"{st['Ma']:7.3f} "
            f"{st['alpha']:12.2f} "
            f"{st['beta']:11.2f} "
            f"{1e3*st['r']:9.2f} "
            f"{1e3*st['H']:9.2f}"
        )
    print("-" * len(header))
    print()


def stage_performance_table(out):
    perf = out["stage_performance"]

    rows = [
        ("Total-to-total efficiency", perf["efficiency_tt"], "-"),
        ("Total-to-static efficiency", perf["efficiency_ts"], "-"),
        ("Pressure ratio (t–s)", perf["pressure_ratio_ts"], "-"),
        ("Volume ratio (t–s)", perf["volume_ratio_ts"], "-"),
        ("Flow coefficient", perf["flow_coefficient"], "-"),
        ("Work coefficient", perf["work_coefficient"], "-"),
        ("Degree of reaction", perf["degree_reaction"], "-"),
        ("Blade velocity ratio", perf["blade_velocity_ratio"], "-"),
        ("Specific speed", perf["specific_speed"], "-"),
        ("Spouting velocity", perf["spouting_velocity"], "m/s"),
        ("Rotor exit blade speed", perf["rotor_exit_velocity"], "m/s"),
        ("Rotational speed", perf["rotational_speed"], "rpm"),
        ("Isentropic power", perf["power_isentropic"] / 1e3, "kW"),
        ("Actual power (t–t)", perf["power_actual_tt"] / 1e3, "kW"),
        ("Actual power (t–s)", perf["power_actual_ts"] / 1e3, "kW"),
    ]

    return [
        {"Quantity": name, "Value": float(value), "Unit": unit}
        for name, value, unit in rows
    ]


def geometry_table(out):
    stator = out["geometry"]["stator"]
    rotor = out["geometry"]["rotor"]

    # Sanity check: geometry dictionaries must match
    if stator.keys() != rotor.keys():
        raise ValueError("Stator and rotor geometry keys do not match")

    def pretty_name(key):
        return key.replace("_", " ").capitalize()

    def infer_unit_and_scale(key):
        if key in {
            "radius_in",
            "radius_out",
            "height",
            "chord",
            "spacing",
            "opening",
            "maximum_thickness",
            "leading_edge_radius",
            "trailing_edge_thickness",
        }:
            return "mm", 1e3

        if key in {
            "metal_angle_in",
            "metal_angle_out",
            "flaring_angle",
            "trailing_edge_wedge_angle",
        }:
            return "deg", 1.0

        if key in {
            "solidity",
            "aspect_ratio",
            "maximum_thickness_location",
        }:
            return "-", 1.0

        if key == "blade_count":
            return "-", 1.0

        # Fallback
        return "-", 1.0

    rows = []
    for key in stator.keys():
        unit, scale = infer_unit_and_scale(key)

        rows.append(
            {
                "Variable": pretty_name(key),
                "Stator": float(stator[key] * scale),
                "Rotor": float(rotor[key] * scale),
                "Unit": unit,
            }
        )

    return rows


def flow_stations_table(out):
    rows = []

    for i, st in enumerate(out["flow_stations"]):
        rows.append(
            {
                "Station": i,
                "p [bar]": st["p"] / 1e5,
                "T [K]": st["T"],
                "ρ [kg/m³]": st["d"],
                "q [-]": st["q"],
                "v [m/s]": st["v"],
                "w [m/s]": st["w"],
                "u [m/s]": st["u"],
                "Ma [-]": st["Ma"],
                "α [deg]": st["alpha"],
                "β [deg]": st["beta"],
                "r [mm]": 1e3 * st["r"],
                "H [mm]": 1e3 * st["H"],
            }
        )

    return rows

def generate_meanline_report(
    out_left,
    out_right=None,
    left_name="Case A",
    right_name="Case B",
    filename="meanline_report.docx",
):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, ns

    doc = Document()

    # ---------------------------------
    # Helpers
    # ---------------------------------
    def prettify_name(key):
        return key.replace("_", " ").capitalize()

    def format_value(v, scale=1.0):
        if isinstance(v, (int, float)):
            return f"{v * scale:.3f}"
        return "-" if v is None else str(v)

    def set_cell_text(cell, text, bold=False, italic=False, align="center"):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if align == "left" else WD_ALIGN_PARAGRAPH.CENTER
        )
        run = p.add_run(str(text))
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    def set_cell_borders(cell, top=False, bottom=False):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        borders = OxmlElement("w:tcBorders")

        def _border(tag):
            el = OxmlElement(tag)
            el.set(ns.qn("w:val"), "single")
            el.set(ns.qn("w:sz"), "8")
            el.set(ns.qn("w:space"), "0")
            el.set(ns.qn("w:color"), "000000")
            return el

        if top:
            borders.append(_border("w:top"))
        if bottom:
            borders.append(_border("w:bottom"))

        tcPr.append(borders)

    # ---------------------------------
    # Create table
    # ---------------------------------
    ncols = 4 if out_right else 3
    table = doc.add_table(rows=1, cols=ncols)
    table.autofit = True

    # Header
    hdr_row = table.rows[0]
    hdr = hdr_row.cells
    set_cell_text(hdr[0], "Variable", bold=True, align="left")
    set_cell_text(hdr[1], left_name, bold=True)
    if out_right:
        set_cell_text(hdr[2], right_name, bold=True)
        set_cell_text(hdr[3], "Unit", bold=True)
    else:
        set_cell_text(hdr[2], "Unit", bold=True)

    # set_row_borders(hdr_row, top=True, bottom=True)
    for cell in hdr_row.cells:
        set_cell_borders(cell, top=True, bottom=True)


    # ---------------------------------
    # Row writers
    # ---------------------------------
    def add_section(title):
        row = table.add_row()
        cells = row.cells
        set_cell_text(cells[0], title, bold=True, italic=True, align="left")
        for c in cells[1:]:
            set_cell_text(c, "")
        
        for cell in row.cells:
            set_cell_borders(cell, top=True, bottom=True)


        # set_row_borders(row, top=True, bottom=True)

    def add_row(name, v1, v2, unit, scale=1.0):
        row = table.add_row().cells
        set_cell_text(row[0], prettify_name(name), align="left")
        set_cell_text(row[1], format_value(v1, scale))
        if out_right:
            set_cell_text(row[2], format_value(v2, scale))
            set_cell_text(row[3], unit)
        else:
            set_cell_text(row[2], unit)

    # ---------------------------------
    # Operating conditions
    # ---------------------------------
    add_section("Operating conditions")
    oc_units = {
        "fluid": ("-", 1.0),
        "stage_type": ("-", 1.0),
        "inlet_property_pair": ("-", 1.0),
        "inlet_property_1": ("bar", 1e-5),
        "inlet_property_2": ("-", 1.0),
        "exit_pressure": ("bar", 1e-5),
        "mass_flow_rate": ("kg/s", 1.0),
    }

    for k, (unit, scale) in oc_units.items():
        add_row(
            k,
            out_left["inputs"].get(k),
            out_right["inputs"].get(k) if out_right else None,
            unit,
            scale,
        )

    # ---------------------------------
    # Design variables
    # ---------------------------------
    add_section("Design variables")
    dv_units = {
        "stator_inlet_angle": ("deg", 1.0),
        "stator_exit_angle": ("deg", 1.0),
        "blade_velocity_ratio": ("-", 1.0),
        "degree_reaction": ("-", 1.0),
        "radius_ratio_12": ("-", 1.0),
        "radius_ratio_23": ("-", 1.0),
        "radius_ratio_34": ("-", 1.0),
        "height_radius_ratio": ("-", 1.0),
        "zweiffel_stator": ("-", 1.0),
        "zweiffel_rotor": ("-", 1.0),
        "loss_coeff_stator": ("-", 1.0),
        "loss_coeff_rotor": ("-", 1.0),
    }

    for k, (unit, scale) in dv_units.items():
        add_row(
            k,
            out_left["inputs"].get(k),
            out_right["inputs"].get(k) if out_right else None,
            unit,
            scale,
        )

    # ---------------------------------
    # Stage performance
    # ---------------------------------
    add_section("Stage performance")
    sp_units = {
        "efficiency_tt": ("-", 1.0),
        "efficiency_ts": ("-", 1.0),
        "pressure_ratio_ts": ("-", 1.0),
        "volume_ratio_ts": ("-", 1.0),
        "flow_coefficient": ("-", 1.0),
        "work_coefficient": ("-", 1.0),
        "specific_speed": ("-", 1.0),
        "rotational_speed": ("rpm", 1.0),
        "spouting_velocity": ("m/s", 1.0),
        "power_isentropic": ("W", 1.0),
        "power_actual_tt": ("W", 1.0),
        "power_actual_ts": ("W", 1.0),
    }

    for k, (unit, scale) in sp_units.items():
        add_row(
            k,
            out_left["stage_performance"].get(k),
            out_right["stage_performance"].get(k) if out_right else None,
            unit,
            scale,
        )

    # ---------------------------------
    # Geometry
    # ---------------------------------
    for part in ("stator", "rotor"):
        add_section(f"{part.capitalize()} geometry")
        geo_units = {
            "blade_count": ("-", 1.0),
            "radius_in": ("mm", 1e3),
            "radius_out": ("mm", 1e3),
            "height": ("mm", 1e3),
            "chord": ("mm", 1e3),
            "spacing": ("mm", 1e3),
            "opening": ("mm", 1e3),
            "solidity": ("-", 1.0),
            "aspect_ratio": ("-", 1.0),
            "flaring_angle": ("deg", 1.0),
            "metal_angle_in": ("deg", 1.0),
            "metal_angle_out": ("deg", 1.0),
        }

        for k, (unit, scale) in geo_units.items():
            add_row(
                k,
                out_left["geometry"][part].get(k),
                out_right["geometry"][part].get(k) if out_right else None,
                unit,
                scale,
            )

    # Bottom rule for entire table (last row)
    last_row = table.rows[-1]
    for cell in last_row.cells:
        set_cell_borders(cell, bottom=True)
        
    doc.save(filename)
    return filename
